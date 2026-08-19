# =================================================================
# IMPORTS AND ENVIRONMENT
# =================================================================

import os
import json
from pydantic import BaseModel, Field
from groq import Groq
from openai import OpenAI  # Used for local LLM support (LM Studio / Ollama)
from dotenv import load_dotenv
from pathlib import Path

load_dotenv()
my_api_key = os.getenv("GROQ_API_KEY")
if not my_api_key:
    raise ValueError("GROQ_API_KEY is missing — add it to your .env file")


# =================================================================
# LLM CLIENT CONFIGURATION
# =================================================================

# --- Groq (cloud) ---
client_groq = Groq(api_key=my_api_key)
groq_model = "openai/gpt-oss-120b"

# --- Local LLM (LM Studio / Ollama) ---
# LM Studio:  base_url="http://localhost:1234/v1"
# Ollama:     base_url="http://localhost:11434/v1"
client_local = OpenAI(
    base_url="http://localhost:1234/v1",
    api_key="anything"              # local servers don't need a real key
)
local_model = "google/gemma-4-e2b"

# =================================================================
# TOGGLE — flip to True to use local LLM instead of Groq
USE_LOCAL_LLM = False
# =================================================================


# =================================================================
# DOCUMENT READERS
# =================================================================

import pdfplumber
from docx import Document


def read_pdf(file_path):
    text = ""
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()

            if page_text:
                text += page_text + "\n"

    return text


def read_docx(file_path):
    document = Document(file_path)
    text = ""
    # Read paragraphs
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            text += paragraph.text + "\n"

    # Read tables
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text += cell.text + "\n"

    return text


def read_txt(file_path):
    with open(file_path, "r", encoding="utf-8") as file:
        return file.read()


def read_document(file_path):
    if file_path.suffix.lower() == ".pdf":
        return read_pdf(file_path)

    elif file_path.suffix.lower() == ".docx":
        return read_docx(file_path)

    elif file_path.suffix.lower() == ".txt":
        return read_txt(file_path)

    else:
        print(f"Unsupported file format: {file_path.name}")


# =================================================================
# PYDANTIC MODELS
# =================================================================

class CandidateProfile(BaseModel):
    name: str | None = None
    skills: list[str] = Field(default_factory=list)
    experience_years: float | None = None
    education: list[str] = Field(default_factory=list)
    projects: list[str] = Field(default_factory=list)
    certifications: list[str] = Field(default_factory=list)


class RoleRequirements(BaseModel):
    role_name: str | None = None
    required_skills: list[str] = Field(default_factory=list)
    preferred_skills: list[str] = Field(default_factory=list)
    min_experience: float | None = None
    responsibilities: list[str] = Field(default_factory=list)


# =================================================================
# LLM EXTRACTION FUNCTIONS
# =================================================================

resume_schema = CandidateProfile.model_json_schema()


def parse_resume(resume_text):
    system_prompt = f"""
You are a professional resume information extraction assistant.

Extract a structured candidate profile from the resume text provided by the
user.

Rules:
1. Extract only facts explicitly stated in the resume.
2. Do not invent, guess, or infer missing information.
3. Keep skills, education, projects, and certifications as concise strings.
4. Convert experience to years only when the resume clearly states it.
5. Use null for missing optional values and [] for missing list values.
6. Return exactly one valid JSON object with no Markdown or explanation.
7. Follow this exact schema:

{resume_schema}
"""
    user_prompt = f"""
Extract the candidate profile from the resume below.

<resume_text>
{resume_text}
</resume_text>

Return one JSON object matching the schema from the system instructions.
"""

    message_system = {"role": "system", "content": system_prompt}
    message_user = {"role": "user", "content": user_prompt}
    messages = [message_system, message_user]
    response_format = {"type": "json_object"}

    # Pick client and model based on toggle
    client = client_local if USE_LOCAL_LLM else client_groq
    model = local_model if USE_LOCAL_LLM else groq_model

    response = client.chat.completions.create(
        model=model,
        messages=messages,
        response_format=response_format,
        temperature=0
    )

    answer = response.choices[0].message.content
    raw_json_output = answer  # raw JSON string from LLM
    data = json.loads(raw_json_output)
    resume_data_object = CandidateProfile(**data)
    return resume_data_object


role_schema = RoleRequirements.model_json_schema()


def parse_job_description(jd_text):
    system_prompt = f"""
You are a professional job-description analysis assistant.

Extract structured role requirements from the job description provided by the
user.

Rules:
1. Extract only facts explicitly stated in the job description.
2. Put mandatory, essential, or required qualifications in required_skills.
3. Put bonus, desirable, or explicitly preferred qualifications in
   preferred_skills.
4. Never classify a preferred skill as required.
5. Extract min_experience only when a minimum is clearly stated.
6. Write responsibilities as concise factual strings.
7. Do not invent requirements, responsibilities, or experience values.
8. Use null for missing optional values and [] for missing list values.
9. Return exactly one valid JSON object with no Markdown or explanation.
10. Follow this exact schema:

{role_schema}
"""
    user_prompt = f"""
Extract the role requirements from the job description below.

<job_description>
{jd_text}
</job_description>

Return one JSON object matching the schema from the system instructions.
"""

    message_system = {"role": "system", "content": system_prompt}
    message_user = {"role": "user", "content": user_prompt}
    messages = [message_system, message_user]
    response_format = {"type": "json_object"}

    # Pick client and model based on toggle
    client = client_local if USE_LOCAL_LLM else client_groq
    model = local_model if USE_LOCAL_LLM else groq_model

    response = client.chat.completions.create(
        model=model,
        messages=messages,
        response_format=response_format,
        temperature=0
    )

    answer = response.choices[0].message.content
    raw_json_output = answer
    data = json.loads(raw_json_output)
    role_data_object = RoleRequirements(**data)
    return role_data_object
